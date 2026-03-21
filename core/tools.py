"""
Permafrost Tool System — Let AI execute tools to interact with the system.

Built-in tools:
  - bash: Execute shell commands
  - read_file: Read file contents
  - write_file: Write/create files
  - list_files: List directory contents
  - web_search: Search the web (placeholder)

Tools are defined as functions with schema (for AI function calling).
Security layer (security.py) controls which tools are allowed.
"""

import json
import logging
import os
import re
import subprocess

log = logging.getLogger("permafrost.tools")

# ── Tool Registry ─────────────────────────────────────────────

TOOLS: dict[str, dict] = {}


def register_tool(name: str, description: str, parameters: dict):
    """Decorator to register a tool with its schema."""
    def decorator(func):
        TOOLS[name] = {
            "function": func,
            "description": description,
            "parameters": parameters,
        }
        return func
    return decorator


def get_tools_schema(provider_type: str = "openai") -> list[dict]:
    """Convert PF tools to native API function calling schema.

    Args:
        provider_type: "openai" (GPT/OpenRouter), "claude" (Anthropic), "gemini" (Google)

    Returns list of tool definitions in the provider's native schema.
    """
    schemas = []
    gemini_funcs = []  # Gemini needs all functions in ONE object

    for name, info in TOOLS.items():
        # Build JSON Schema properties from our simple parameter format
        properties = {}
        required = []
        for param_name, param_info in info["parameters"].items():
            ptype = param_info.get("type", "string")
            # Gemini doesn't support "number" type — use "integer" or "string"
            if provider_type == "gemini" and ptype == "number":
                ptype = "string"
            prop = {"type": ptype}
            if "description" in param_info:
                prop["description"] = param_info["description"]
            properties[param_name] = prop
            if param_info.get("required", False):
                required.append(param_name)

        if provider_type == "openai":
            schemas.append({
                "type": "function",
                "function": {
                    "name": name,
                    "description": info["description"],
                    "parameters": {
                        "type": "object",
                        "properties": properties,
                        "required": required,
                    },
                },
            })
        elif provider_type == "claude":
            schemas.append({
                "name": name,
                "description": info["description"],
                "input_schema": {
                    "type": "object",
                    "properties": properties,
                    "required": required,
                },
            })
        elif provider_type == "gemini":
            func_decl = {
                "name": name,
                "description": info["description"],
            }
            # Only add parameters if there are any (Gemini errors on empty)
            if properties:
                params = {"type": "object", "properties": properties}
                if required:
                    params["required"] = required
                func_decl["parameters"] = params
            gemini_funcs.append(func_decl)

    # Gemini: wrap all functions in a single tool object
    if provider_type == "gemini" and gemini_funcs:
        schemas = [{"function_declarations": gemini_funcs}]

    return schemas


# ── Built-in Tools ────────────────────────────────────────────

@register_tool(
    "bash",
    "Execute a shell command",
    {"command": {"type": "string", "description": "Shell command to run"}},
)
def tool_bash(command: str, **kwargs) -> str:
    """Execute a shell command and return stdout + stderr."""
    try:
        result = subprocess.run(
            command, shell=True, capture_output=True, text=True,
            timeout=30, encoding="utf-8", errors="replace",
        )
        output = result.stdout
        if result.stderr:
            output += f"\n[stderr] {result.stderr}"
        return output[:4000]  # cap output
    except subprocess.TimeoutExpired:
        return "[error] Command timed out (30s)"
    except Exception as e:
        return f"[error] {e}"


@register_tool(
    "read_file",
    "Read a file's contents",
    {"path": {"type": "string", "description": "File path to read"}},
)
def tool_read_file(path: str, **kwargs) -> str:
    """Read file contents (capped at 10K chars)."""
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read(10000)
        return content
    except Exception as e:
        return f"[error] {e}"


@register_tool(
    "write_file",
    "Write content to a file (creates parent directories if needed)",
    {
        "path": {"type": "string", "description": "File path to write"},
        "content": {"type": "string", "description": "Content to write"},
    },
)
def tool_write_file(path: str, content: str, **kwargs) -> str:
    """Write content to a file, creating directories as needed."""
    try:
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"Written {len(content)} chars to {path}"
    except Exception as e:
        return f"[error] {e}"


@register_tool(
    "list_files",
    "List files in a directory",
    {"path": {"type": "string", "description": "Directory path (default: current dir)"}},
)
def tool_list_files(path: str = ".", **kwargs) -> str:
    """List directory entries (capped at 100)."""
    try:
        entries = os.listdir(path)
        return "\n".join(entries[:100])
    except Exception as e:
        return f"[error] {e}"


@register_tool(
    "python_exec",
    "Execute Python code",
    {"code": {"type": "string", "description": "Python code to execute"}},
)
def tool_python_exec(code: str, **kwargs) -> str:
    """Execute a Python code snippet in a subprocess."""
    try:
        result = subprocess.run(
            ["python", "-c", code], capture_output=True, text=True,
            timeout=30, encoding="utf-8", errors="replace",
        )
        output = result.stdout
        if result.stderr:
            output += f"\n[stderr] {result.stderr}"
        return output[:4000]
    except subprocess.TimeoutExpired:
        return "[error] Timeout (30s)"
    except Exception as e:
        return f"[error] {e}"


@register_tool(
    "web_fetch",
    "Fetch content from a URL",
    {"url": {"type": "string", "description": "URL to fetch"}},
)
def tool_web_fetch(url: str, **kwargs) -> str:
    """Fetch webpage content, stripping HTML tags for readability."""
    import requests
    try:
        r = requests.get(url, timeout=15, headers={"User-Agent": "Permafrost/1.0"})
        text = re.sub(r'<[^>]+>', '', r.text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text[:4000]
    except Exception as e:
        return f"[error] {e}"


@register_tool(
    "search_web",
    "Search the web using DuckDuckGo",
    {"query": {"type": "string", "description": "Search query"}},
)
def tool_search_web(query: str, **kwargs) -> str:
    """Search the web via DuckDuckGo Instant Answer API."""
    import requests
    try:
        r = requests.get(
            "https://api.duckduckgo.com/",
            params={"q": query, "format": "json"},
            timeout=10,
        )
        data = r.json()
        results = []
        if data.get("Abstract"):
            results.append(f"Summary: {data['Abstract']}")
        for topic in data.get("RelatedTopics", [])[:5]:
            if isinstance(topic, dict) and topic.get("Text"):
                results.append(f"- {topic['Text']}")
        return "\n".join(results) if results else "No results found."
    except Exception as e:
        return f"[error] {e}"


@register_tool(
    "edit_file",
    "Find and replace text in a file",
    {
        "path": {"type": "string", "description": "File path"},
        "old_text": {"type": "string", "description": "Text to find"},
        "new_text": {"type": "string", "description": "Replacement text"},
    },
)
def tool_edit_file(path: str, old_text: str, new_text: str, **kwargs) -> str:
    """Find and replace first occurrence of old_text with new_text in a file."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        if old_text not in content:
            return f"[error] old_text not found in {path}"
        content = content.replace(old_text, new_text, 1)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"Replaced in {path}"
    except Exception as e:
        return f"[error] {e}"


# ── Utility Tools ─────────────────────────────────────────────

@register_tool("get_datetime", "Get current date, time, and timezone info", {})
def tool_get_datetime(**kwargs) -> str:
    from datetime import datetime
    now = datetime.now()
    return f"Date: {now.strftime('%Y-%m-%d')}\nTime: {now.strftime('%H:%M:%S')}\nDay: {now.strftime('%A')}\nTimestamp: {now.isoformat()}"


@register_tool("calculate", "Evaluate a math expression safely", {
    "expression": {"type": "string", "description": "Math expression (e.g. '2**10', 'sqrt(144)', '3.14*5**2')"},
})
def tool_calculate(expression: str, **kwargs) -> str:
    import math
    allowed = {
        "abs": abs, "round": round, "min": min, "max": max,
        "sum": sum, "len": len, "int": int, "float": float,
        "sqrt": math.sqrt, "ceil": math.ceil, "floor": math.floor,
        "log": math.log, "log10": math.log10, "log2": math.log2,
        "sin": math.sin, "cos": math.cos, "tan": math.tan,
        "pi": math.pi, "e": math.e, "pow": pow,
    }
    try:
        result = eval(expression, {"__builtins__": {}}, allowed)
        return str(result)
    except Exception as e:
        return f"[error] {e}"


@register_tool("http_request", "Make an HTTP request (GET/POST)", {
    "url": {"type": "string", "description": "URL to request"},
    "method": {"type": "string", "description": "GET or POST (default: GET)"},
    "body": {"type": "string", "description": "POST body (JSON string, optional)"},
    "headers": {"type": "string", "description": "JSON string of headers (optional)"},
})
def tool_http_request(url: str, method: str = "GET", body: str = "", headers: str = "", **kwargs) -> str:
    import requests
    try:
        h = json.loads(headers) if headers else {}
        h.setdefault("User-Agent", "Permafrost/1.0")
        if method.upper() == "POST":
            data = json.loads(body) if body else {}
            r = requests.post(url, json=data, headers=h, timeout=15)
        else:
            r = requests.get(url, headers=h, timeout=15)
        return f"Status: {r.status_code}\n{r.text[:3000]}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("json_read", "Read and parse a JSON file", {
    "path": {"type": "string", "description": "Path to JSON file"},
    "key": {"type": "string", "description": "Optional dot-notation key to extract (e.g. 'data.items')"},
})
def tool_json_read(path: str, key: str = "", **kwargs) -> str:
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if key:
            for k in key.split("."):
                if isinstance(data, dict):
                    data = data.get(k, f"[key '{k}' not found]")
                elif isinstance(data, list) and k.isdigit():
                    data = data[int(k)]
                else:
                    return f"[error] Cannot navigate '{k}' in {type(data).__name__}"
        return json.dumps(data, ensure_ascii=False, indent=2)[:4000]
    except Exception as e:
        return f"[error] {e}"


@register_tool("json_write", "Write data to a JSON file", {
    "path": {"type": "string", "description": "Path to JSON file"},
    "data": {"type": "string", "description": "JSON string to write"},
})
def tool_json_write(path: str, data: str, **kwargs) -> str:
    try:
        parsed = json.loads(data)
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(parsed, f, ensure_ascii=False, indent=2)
        return f"Written JSON to {path}"
    except json.JSONDecodeError as e:
        return f"[error] Invalid JSON: {e}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("send_notification", "Send a message to the user through all enabled channels", {
    "message": {"type": "string", "description": "Message to send"},
})
def tool_send_notification(message: str, **kwargs) -> str:
    try:
        from core.scheduler import PFScheduler
        sched = PFScheduler()
        sched.notify_user(message)
        return f"Notification sent: {message[:100]}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("append_file", "Append text to the end of a file", {
    "path": {"type": "string", "description": "File path"},
    "content": {"type": "string", "description": "Content to append"},
})
def tool_append_file(path: str, content: str, **kwargs) -> str:
    try:
        with open(path, "a", encoding="utf-8") as f:
            f.write(content)
        return f"Appended {len(content)} chars to {path}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("grep_files", "Search for a pattern in files (recursive)", {
    "pattern": {"type": "string", "description": "Text or regex pattern to search"},
    "path": {"type": "string", "description": "Directory to search (default: current dir)"},
    "file_pattern": {"type": "string", "description": "File glob pattern (e.g. '*.py', '*.md')"},
})
def tool_grep_files(pattern: str, path: str = ".", file_pattern: str = "*", **kwargs) -> str:
    import glob
    try:
        matches = []
        for filepath in glob.glob(os.path.join(path, "**", file_pattern), recursive=True):
            if os.path.isfile(filepath):
                try:
                    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                        for i, line in enumerate(f, 1):
                            if pattern.lower() in line.lower():
                                matches.append(f"{filepath}:{i}: {line.strip()[:100]}")
                                if len(matches) >= 20:
                                    return "\n".join(matches) + "\n... (truncated)"
                except (OSError, UnicodeDecodeError):
                    pass
        return "\n".join(matches) if matches else "No matches found."
    except Exception as e:
        return f"[error] {e}"


# ── System & OS Tools ─────────────────────────────────────────

@register_tool("download_file", "Download a file from a URL to local disk", {
    "url": {"type": "string", "description": "URL to download"},
    "path": {"type": "string", "description": "Local path to save the file"},
})
def tool_download_file(url: str, path: str, **kwargs) -> str:
    import requests
    try:
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        r = requests.get(url, stream=True, timeout=30, headers={"User-Agent": "Permafrost/1.0"})
        r.raise_for_status()
        total = 0
        with open(path, "wb") as f:
            for chunk in r.iter_content(chunk_size=8192):
                f.write(chunk)
                total += len(chunk)
        size_kb = total / 1024
        return f"Downloaded {size_kb:.1f} KB to {path}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("system_info", "Get system information (CPU, memory, disk, OS)", {})
def tool_system_info(**kwargs) -> str:
    import platform
    lines = [
        f"OS: {platform.system()} {platform.release()} ({platform.machine()})",
        f"Python: {platform.python_version()}",
        f"Hostname: {platform.node()}",
    ]
    try:
        import psutil
        mem = psutil.virtual_memory()
        disk = psutil.disk_usage("/")
        lines.append(f"CPU: {psutil.cpu_count()} cores, {psutil.cpu_percent()}% used")
        lines.append(f"RAM: {mem.used/1e9:.1f}/{mem.total/1e9:.1f} GB ({mem.percent}%)")
        lines.append(f"Disk: {disk.used/1e9:.1f}/{disk.total/1e9:.1f} GB ({disk.percent}%)")
    except ImportError:
        lines.append("(install psutil for CPU/RAM/disk details)")
    return "\n".join(lines)


@register_tool("process_list", "List running processes (optional filter)", {
    "filter": {"type": "string", "description": "Filter by process name (optional)"},
})
def tool_process_list(filter: str = "", **kwargs) -> str:
    try:
        import psutil
        procs = []
        for p in psutil.process_iter(["pid", "name", "cpu_percent", "memory_info"]):
            try:
                info = p.info
                if filter and filter.lower() not in info["name"].lower():
                    continue
                mem_mb = info["memory_info"].rss / 1e6 if info.get("memory_info") else 0
                procs.append(f"  PID {info['pid']:6d} | {info['name'][:25]:25s} | {mem_mb:.0f} MB")
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                pass
        procs.sort()
        return f"Processes ({len(procs)}):\n" + "\n".join(procs[:30])
    except ImportError:
        return tool_bash(command='tasklist /FO TABLE /NH' if os.name == 'nt' else 'ps aux --sort=-rss | head -20')


@register_tool("kill_process", "Kill a process by PID", {
    "pid": {"type": "number", "description": "Process ID to kill"},
})
def tool_kill_process(pid: int, **kwargs) -> str:
    try:
        import psutil
        p = psutil.Process(int(pid))
        name = p.name()
        p.terminate()
        return f"Terminated: {name} (PID {pid})"
    except ImportError:
        return tool_bash(command=f'taskkill /PID {int(pid)} /F' if os.name == 'nt' else f'kill {int(pid)}')
    except Exception as e:
        return f"[error] {e}"


@register_tool("compress", "Compress files/directories into a zip archive", {
    "path": {"type": "string", "description": "File or directory to compress"},
    "output": {"type": "string", "description": "Output zip file path"},
})
def tool_compress(path: str, output: str, **kwargs) -> str:
    import zipfile
    try:
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zf:
            if os.path.isdir(path):
                for root, dirs, files in os.walk(path):
                    for f in files:
                        fp = os.path.join(root, f)
                        arc = os.path.relpath(fp, os.path.dirname(path))
                        zf.write(fp, arc)
            else:
                zf.write(path, os.path.basename(path))
        size_kb = os.path.getsize(output) / 1024
        return f"Compressed to {output} ({size_kb:.1f} KB)"
    except Exception as e:
        return f"[error] {e}"


@register_tool("extract", "Extract a zip/tar archive", {
    "path": {"type": "string", "description": "Archive file path (.zip, .tar.gz, .tar)"},
    "output": {"type": "string", "description": "Directory to extract to"},
})
def tool_extract(path: str, output: str, **kwargs) -> str:
    try:
        os.makedirs(output, exist_ok=True)
        if path.endswith(".zip"):
            import zipfile
            with zipfile.ZipFile(path, "r") as zf:
                zf.extractall(output)
                return f"Extracted {len(zf.namelist())} files to {output}"
        elif path.endswith((".tar.gz", ".tgz", ".tar")):
            import tarfile
            with tarfile.open(path, "r:*") as tf:
                tf.extractall(output)
                return f"Extracted to {output}"
        return f"[error] Unsupported archive type: {path}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("diff_files", "Compare two files and show differences", {
    "file1": {"type": "string", "description": "First file path"},
    "file2": {"type": "string", "description": "Second file path"},
})
def tool_diff_files(file1: str, file2: str, **kwargs) -> str:
    import difflib
    try:
        with open(file1, "r", encoding="utf-8", errors="replace") as f:
            lines1 = f.readlines()
        with open(file2, "r", encoding="utf-8", errors="replace") as f:
            lines2 = f.readlines()
        diff = list(difflib.unified_diff(lines1, lines2, fromfile=file1, tofile=file2, lineterm=""))
        if not diff:
            return "Files are identical."
        return "\n".join(diff[:100])
    except Exception as e:
        return f"[error] {e}"


@register_tool("file_hash", "Calculate file hash (MD5, SHA256)", {
    "path": {"type": "string", "description": "File path"},
    "algorithm": {"type": "string", "description": "md5 or sha256 (default: sha256)"},
})
def tool_file_hash(path: str, algorithm: str = "sha256", **kwargs) -> str:
    import hashlib
    try:
        h = hashlib.new(algorithm)
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        size_kb = os.path.getsize(path) / 1024
        return f"{algorithm.upper()}: {h.hexdigest()}\nSize: {size_kb:.1f} KB"
    except Exception as e:
        return f"[error] {e}"


@register_tool("clipboard", "Read or write system clipboard", {
    "action": {"type": "string", "description": "read or write"},
    "text": {"type": "string", "description": "Text to write (only for write action)"},
})
def tool_clipboard(action: str = "read", text: str = "", **kwargs) -> str:
    try:
        if action == "write" and text:
            if os.name == "nt":
                process = subprocess.Popen(["clip"], stdin=subprocess.PIPE)
                process.communicate(text.encode("utf-16-le"))
            else:
                process = subprocess.Popen(["xclip", "-selection", "clipboard"], stdin=subprocess.PIPE)
                process.communicate(text.encode("utf-8"))
            return f"Copied {len(text)} chars to clipboard"
        else:
            if os.name == "nt":
                result = subprocess.run(["powershell", "-command", "Get-Clipboard"], capture_output=True, text=True, timeout=5)
                return result.stdout.strip() or "[clipboard empty]"
            else:
                result = subprocess.run(["xclip", "-selection", "clipboard", "-o"], capture_output=True, text=True, timeout=5)
                return result.stdout.strip() or "[clipboard empty]"
    except Exception as e:
        return f"[error] {e}"


@register_tool("screenshot", "Take a screenshot of the screen", {
    "path": {"type": "string", "description": "Output image path (default: screenshot.png)"},
})
def tool_screenshot(path: str = "screenshot.png", **kwargs) -> str:
    try:
        from PIL import ImageGrab
        img = ImageGrab.grab()
        img.save(path)
        return f"Screenshot saved: {path} ({img.size[0]}x{img.size[1]})"
    except ImportError:
        try:
            if os.name == "nt":
                result = subprocess.run(
                    ["powershell", "-command",
                     f"Add-Type -AssemblyName System.Windows.Forms; "
                     f"[System.Windows.Forms.Screen]::PrimaryScreen | ForEach-Object {{ "
                     f"$b = New-Object System.Drawing.Bitmap($_.Bounds.Width, $_.Bounds.Height); "
                     f"$g = [System.Drawing.Graphics]::FromImage($b); "
                     f"$g.CopyFromScreen($_.Bounds.Location, [System.Drawing.Point]::Empty, $_.Bounds.Size); "
                     f"$b.Save('{path}') }}"],
                    capture_output=True, text=True, timeout=10,
                )
                if os.path.exists(path):
                    return f"Screenshot saved: {path}"
                return f"[error] Screenshot failed: {result.stderr}"
            return "[error] Install Pillow: pip install Pillow"
        except Exception as e:
            return f"[error] {e}"
    except Exception as e:
        return f"[error] {e}"


# ── Text & Data Tools ─────────────────────────────────────────

@register_tool("encode_decode", "Encode/decode text (base64, URL, HTML entities)", {
    "text": {"type": "string", "description": "Text to encode/decode"},
    "method": {"type": "string", "description": "base64_encode, base64_decode, url_encode, url_decode, html_escape, html_unescape"},
})
def tool_encode_decode(text: str, method: str = "base64_encode", **kwargs) -> str:
    import base64, urllib.parse, html
    try:
        if method == "base64_encode":
            return base64.b64encode(text.encode()).decode()
        elif method == "base64_decode":
            return base64.b64decode(text).decode()
        elif method == "url_encode":
            return urllib.parse.quote(text, safe="")
        elif method == "url_decode":
            return urllib.parse.unquote(text)
        elif method == "html_escape":
            return html.escape(text)
        elif method == "html_unescape":
            return html.unescape(text)
        return f"[error] Unknown method: {method}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("regex_extract", "Extract data from text using regex", {
    "text": {"type": "string", "description": "Text to search"},
    "pattern": {"type": "string", "description": "Regex pattern (use groups to capture)"},
})
def tool_regex_extract(text: str, pattern: str, **kwargs) -> str:
    try:
        matches = re.findall(pattern, text)
        if not matches:
            return "No matches."
        return json.dumps(matches[:20], ensure_ascii=False)
    except re.error as e:
        return f"[error] Invalid regex: {e}"


@register_tool("text_stats", "Count words, characters, lines in text or file", {
    "text": {"type": "string", "description": "Text to analyze (or file path if starts with @)"},
})
def tool_text_stats(text: str, **kwargs) -> str:
    try:
        if text.startswith("@") and os.path.isfile(text[1:]):
            with open(text[1:], "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        lines = text.count("\n") + 1
        words = len(text.split())
        chars = len(text)
        return f"Lines: {lines}\nWords: {words}\nChars: {chars}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("generate_password", "Generate a secure random password", {
    "length": {"type": "number", "description": "Password length (default: 16)"},
    "no_symbols": {"type": "string", "description": "Set to 'true' for alphanumeric only"},
})
def tool_generate_password(length: int = 16, no_symbols: str = "false", **kwargs) -> str:
    import secrets, string
    chars = string.ascii_letters + string.digits
    if no_symbols.lower() != "true":
        chars += "!@#$%&*-_=+"
    pwd = "".join(secrets.choice(chars) for _ in range(int(length)))
    return pwd


@register_tool("generate_uuid", "Generate a UUID", {})
def tool_generate_uuid(**kwargs) -> str:
    import uuid
    return str(uuid.uuid4())


# ── Network Tools ────────────────────────────────────────────

@register_tool("ping", "Check if a host is reachable", {
    "host": {"type": "string", "description": "Hostname or IP address"},
})
def tool_ping(host: str, **kwargs) -> str:
    try:
        flag = "-n" if os.name == "nt" else "-c"
        result = subprocess.run(
            ["ping", flag, "3", host],
            capture_output=True, text=True, timeout=15,
        )
        return result.stdout[:2000] if result.returncode == 0 else f"Host unreachable:\n{result.stdout[:500]}"
    except subprocess.TimeoutExpired:
        return f"[timeout] {host} did not respond in 15s"
    except Exception as e:
        return f"[error] {e}"


@register_tool("port_check", "Check if a TCP port is open", {
    "host": {"type": "string", "description": "Hostname or IP"},
    "port": {"type": "number", "description": "Port number"},
})
def tool_port_check(host: str, port: int, **kwargs) -> str:
    import socket
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        s.settimeout(5)
        result = s.connect_ex((host, int(port)))
        s.close()
        return f"{host}:{int(port)} — {'OPEN' if result == 0 else 'CLOSED'}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("dns_lookup", "DNS lookup for a domain", {
    "domain": {"type": "string", "description": "Domain to lookup"},
})
def tool_dns_lookup(domain: str, **kwargs) -> str:
    import socket
    try:
        results = socket.getaddrinfo(domain, None)
        ips = list(set(r[4][0] for r in results))
        return f"{domain}:\n" + "\n".join(f"  {ip}" for ip in ips[:10])
    except socket.gaierror as e:
        return f"[error] DNS lookup failed: {e}"


# ── Git Tools ─────────────────────────────────────────────────

@register_tool("git_status", "Show git repository status", {
    "path": {"type": "string", "description": "Repository path (default: current dir)"},
})
def tool_git_status(path: str = ".", **kwargs) -> str:
    try:
        r = subprocess.run(["git", "status", "--short"], capture_output=True,
                          text=True, cwd=path, timeout=10)
        branch = subprocess.run(["git", "branch", "--show-current"], capture_output=True,
                               text=True, cwd=path, timeout=5)
        return f"Branch: {branch.stdout.strip()}\n{r.stdout[:2000] or 'Clean'}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("git_log", "Show recent git commits", {
    "path": {"type": "string", "description": "Repository path (default: current dir)"},
    "count": {"type": "number", "description": "Number of commits (default: 10)"},
})
def tool_git_log(path: str = ".", count: int = 10, **kwargs) -> str:
    try:
        r = subprocess.run(
            ["git", "log", f"--oneline", f"-{int(count)}"],
            capture_output=True, text=True, cwd=path, timeout=10,
        )
        return r.stdout[:3000] or "No commits."
    except Exception as e:
        return f"[error] {e}"


@register_tool("git_diff", "Show git changes (staged or unstaged)", {
    "path": {"type": "string", "description": "Repository path (default: current dir)"},
    "staged": {"type": "string", "description": "Set to 'true' for staged changes only"},
})
def tool_git_diff(path: str = ".", staged: str = "false", **kwargs) -> str:
    try:
        cmd = ["git", "diff"]
        if staged.lower() == "true":
            cmd.append("--staged")
        r = subprocess.run(cmd, capture_output=True, text=True, cwd=path, timeout=10)
        return r.stdout[:4000] or "No changes."
    except Exception as e:
        return f"[error] {e}"


# ── QR Code ──────────────────────────────────────────────────

@register_tool("qrcode_create", "Generate a QR code image", {
    "data": {"type": "string", "description": "Data to encode (URL, text, etc.)"},
    "path": {"type": "string", "description": "Output image path (default: qrcode.png)"},
})
def tool_qrcode_create(data: str, path: str = "qrcode.png", **kwargs) -> str:
    try:
        import qrcode
        img = qrcode.make(data)
        img.save(path)
        return f"QR code saved: {path}"
    except ImportError:
        return "[error] Install qrcode: pip install qrcode[pil]"
    except Exception as e:
        return f"[error] {e}"


# ── PF Schedule Management ───────────────────────────────────

@register_tool("schedule_add", "Add a new scheduled task to PF scheduler", {
    "task_id": {"type": "string", "description": "Unique task ID (snake_case)"},
    "description": {"type": "string", "description": "What the task does"},
    "schedule_type": {"type": "string", "description": "cron, daily, interval, or once"},
    "schedule_value": {"type": "string", "description": "Cron expr, HH:MM, minutes, or ISO datetime"},
    "command": {"type": "string", "description": "Command/instruction for the brain to execute"},
})
def tool_schedule_add(task_id: str, description: str, schedule_type: str,
                      schedule_value: str, command: str, **kwargs) -> str:
    from pathlib import Path
    data_dir = Path(os.path.expanduser("~/.permafrost"))
    schedule_file = data_dir / "schedule.json"
    try:
        schedule = {"tasks": []}
        if schedule_file.exists():
            raw = json.loads(schedule_file.read_text(encoding="utf-8"))
            if isinstance(raw, dict):
                schedule = raw
            elif isinstance(raw, list):
                schedule = {"tasks": raw}

        # Check for duplicate
        for t in schedule["tasks"]:
            if t.get("id") == task_id:
                return f"[error] Task '{task_id}' already exists"

        task = {
            "id": task_id,
            "enabled": True,
            "description": description,
            "command": command,
            "schedule": {},
        }
        if schedule_type == "cron":
            task["schedule"] = {"type": "cron", "cron": schedule_value}
        elif schedule_type == "daily":
            task["schedule"] = {"type": "daily", "time": schedule_value}
        elif schedule_type == "interval":
            task["schedule"] = {"type": "interval", "minutes": int(schedule_value)}
        elif schedule_type == "once":
            task["schedule"] = {"type": "once", "datetime": schedule_value}

        schedule["tasks"].append(task)
        schedule_file.write_text(json.dumps(schedule, indent=2, ensure_ascii=False), encoding="utf-8")
        return f"Task '{task_id}' added ({schedule_type}: {schedule_value})"
    except Exception as e:
        return f"[error] {e}"


@register_tool("schedule_list", "List all PF scheduled tasks", {})
def tool_schedule_list(**kwargs) -> str:
    from pathlib import Path
    schedule_file = Path(os.path.expanduser("~/.permafrost")) / "schedule.json"
    if not schedule_file.exists():
        return "No schedule configured."
    try:
        raw = json.loads(schedule_file.read_text(encoding="utf-8"))
        tasks = raw.get("tasks", raw) if isinstance(raw, dict) else raw
        lines = []
        for t in tasks:
            status = "ON" if t.get("enabled", True) else "OFF"
            sched = t.get("schedule", {})
            stype = sched.get("type", "?")
            sval = sched.get("cron", sched.get("time", sched.get("minutes", sched.get("datetime", "?"))))
            lines.append(f"  [{status}] {t.get('id','?')} ({stype}: {sval}) — {t.get('description','')[:60]}")
        return f"Scheduled tasks ({len(tasks)}):\n" + "\n".join(lines)
    except Exception as e:
        return f"[error] {e}"


# ── Document & Media Tools ────────────────────────────────────

@register_tool("create_pdf", "Create a PDF document from text or HTML content", {
    "path": {"type": "string", "description": "Output PDF file path"},
    "content": {"type": "string", "description": "Text content (or HTML if starts with '<')"},
    "title": {"type": "string", "description": "Document title (optional)"},
})
def tool_create_pdf(path: str, content: str, title: str = "", **kwargs) -> str:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        doc = SimpleDocTemplate(path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        if title:
            story.append(Paragraph(title, styles["Title"]))
            story.append(Spacer(1, 0.5 * cm))
        for line in content.split("\n"):
            if line.strip():
                story.append(Paragraph(line, styles["Normal"]))
                story.append(Spacer(1, 0.2 * cm))
        doc.build(story)
        return f"PDF created: {path}"
    except ImportError:
        # Fallback: use fpdf2 (lighter dependency)
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            # Try to add unicode font
            try:
                pdf.add_font("NotoSans", "", os.path.join(os.path.dirname(__file__), "fonts", "NotoSansCJK-Regular.ttc"), uni=True)
                pdf.set_font("NotoSans", size=12)
            except Exception:
                pdf.set_font("Helvetica", size=12)
            if title:
                pdf.set_font_size(18)
                pdf.cell(0, 10, title, ln=True, align="C")
                pdf.set_font_size(12)
                pdf.ln(5)
            for line in content.split("\n"):
                pdf.multi_cell(0, 7, line)
            pdf.output(path)
            return f"PDF created: {path}"
        except ImportError:
            return "[error] Install reportlab or fpdf2: pip install reportlab fpdf2"


@register_tool("create_spreadsheet", "Create an Excel/CSV spreadsheet", {
    "path": {"type": "string", "description": "Output file path (.xlsx or .csv)"},
    "data": {"type": "string", "description": "JSON array of rows, e.g. [[\"Name\",\"Age\"],[\"Alice\",30]]"},
    "sheet_name": {"type": "string", "description": "Sheet name (xlsx only, default: Sheet1)"},
})
def tool_create_spreadsheet(path: str, data: str, sheet_name: str = "Sheet1", **kwargs) -> str:
    try:
        rows = json.loads(data)
        if not isinstance(rows, list):
            return "[error] Data must be a JSON array of rows"

        if path.endswith(".csv"):
            import csv
            with open(path, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                for row in rows:
                    writer.writerow(row)
            return f"CSV created: {path} ({len(rows)} rows)"
        else:
            try:
                import openpyxl
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = sheet_name
                for row in rows:
                    ws.append(row)
                wb.save(path)
                return f"Excel created: {path} ({len(rows)} rows)"
            except ImportError:
                # Fallback to CSV
                import csv
                csv_path = path.rsplit(".", 1)[0] + ".csv"
                with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
                    writer = csv.writer(f)
                    for row in rows:
                        writer.writerow(row)
                return f"openpyxl not installed, saved as CSV: {csv_path} ({len(rows)} rows)"
    except json.JSONDecodeError as e:
        return f"[error] Invalid JSON data: {e}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("read_spreadsheet", "Read an Excel or CSV file", {
    "path": {"type": "string", "description": "File path (.xlsx, .xls, .csv)"},
    "sheet": {"type": "string", "description": "Sheet name (xlsx only, default: first sheet)"},
    "max_rows": {"type": "number", "description": "Max rows to read (default: 50)"},
})
def tool_read_spreadsheet(path: str, sheet: str = "", max_rows: int = 50, **kwargs) -> str:
    try:
        if path.endswith(".csv"):
            import csv
            with open(path, "r", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                rows = [row for _, row in zip(range(int(max_rows)), reader)]
            return json.dumps(rows, ensure_ascii=False)
        else:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(path, read_only=True)
                ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active
                rows = []
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= int(max_rows):
                        break
                    rows.append([str(c) if c is not None else "" for c in row])
                return json.dumps(rows, ensure_ascii=False)
            except ImportError:
                return "[error] Install openpyxl: pip install openpyxl"
    except Exception as e:
        return f"[error] {e}"


@register_tool("create_document", "Create a Word document (.docx)", {
    "path": {"type": "string", "description": "Output .docx file path"},
    "content": {"type": "string", "description": "Document content (paragraphs separated by newlines)"},
    "title": {"type": "string", "description": "Document title (optional)"},
})
def tool_create_document(path: str, content: str, title: str = "", **kwargs) -> str:
    try:
        from docx import Document
        doc = Document()
        if title:
            doc.add_heading(title, 0)
        for para in content.split("\n"):
            if para.strip():
                if para.startswith("## "):
                    doc.add_heading(para[3:], level=2)
                elif para.startswith("# "):
                    doc.add_heading(para[2:], level=1)
                elif para.startswith("- "):
                    doc.add_paragraph(para[2:], style="List Bullet")
                else:
                    doc.add_paragraph(para)
        doc.save(path)
        return f"Word document created: {path}"
    except ImportError:
        return "[error] Install python-docx: pip install python-docx"
    except Exception as e:
        return f"[error] {e}"


@register_tool("read_pdf", "Read text content from a PDF file", {
    "path": {"type": "string", "description": "PDF file path"},
    "max_pages": {"type": "number", "description": "Max pages to read (default: 20)"},
})
def tool_read_pdf(path: str, max_pages: int = 20, **kwargs) -> str:
    try:
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(path)
            texts = []
            for i, page in enumerate(doc):
                if i >= int(max_pages):
                    break
                texts.append(page.get_text())
            return "\n---\n".join(texts)[:8000]
        except ImportError:
            pass
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            texts = []
            for i, page in enumerate(reader.pages):
                if i >= int(max_pages):
                    break
                texts.append(page.extract_text() or "")
            return "\n---\n".join(texts)[:8000]
        except ImportError:
            return "[error] Install PyMuPDF or pypdf: pip install pymupdf pypdf"
    except Exception as e:
        return f"[error] {e}"


@register_tool("read_image", "Read/describe an image file (returns base64 or OCR text)", {
    "path": {"type": "string", "description": "Image file path"},
    "mode": {"type": "string", "description": "info (default), ocr, or base64"},
})
def tool_read_image(path: str, mode: str = "info", **kwargs) -> str:
    try:
        from PIL import Image
        img = Image.open(path)
        info = f"Size: {img.size[0]}x{img.size[1]}, Mode: {img.mode}, Format: {img.format or 'unknown'}"

        if mode == "base64":
            import base64, io
            buf = io.BytesIO()
            img.save(buf, format=img.format or "PNG")
            b64 = base64.b64encode(buf.getvalue()).decode()
            return f"{info}\nBase64: {b64[:200]}... ({len(b64)} chars total)"
        elif mode == "ocr":
            try:
                import pytesseract
                text = pytesseract.image_to_string(img)
                return f"{info}\nOCR Text:\n{text[:4000]}"
            except ImportError:
                return f"{info}\n[OCR unavailable — install pytesseract]"
        else:
            return info
    except ImportError:
        return "[error] Install Pillow: pip install Pillow"
    except Exception as e:
        return f"[error] {e}"


@register_tool("resize_image", "Resize an image to specified dimensions", {
    "path": {"type": "string", "description": "Input image path"},
    "output": {"type": "string", "description": "Output image path"},
    "width": {"type": "number", "description": "Target width in pixels"},
    "height": {"type": "number", "description": "Target height in pixels (0 = auto aspect ratio)"},
})
def tool_resize_image(path: str, output: str, width: int = 800, height: int = 0, **kwargs) -> str:
    try:
        from PIL import Image
        img = Image.open(path)
        w, h = int(width), int(height)
        if h <= 0:
            ratio = w / img.size[0]
            h = int(img.size[1] * ratio)
        img = img.resize((w, h), Image.LANCZOS)
        img.save(output)
        return f"Resized to {w}x{h}: {output}"
    except ImportError:
        return "[error] Install Pillow: pip install Pillow"
    except Exception as e:
        return f"[error] {e}"


# ── Multi-Agent Tools ─────────────────────────────────────────

@register_tool("spawn_agent", "Create an independent AI agent with its own memory and brain", {
    "name": {"type": "string", "description": "Agent name (lowercase, e.g. 'researcher')"},
    "persona": {"type": "string", "description": "Agent's role/personality (system prompt)"},
    "task": {"type": "string", "description": "Initial task to assign"},
})
def tool_spawn_agent(name: str, persona: str = "", task: str = "", **kwargs) -> str:
    from core.multi_agent import PFMultiAgent
    try:
        ma = PFMultiAgent("main", os.path.expanduser("~/.permafrost"))
        return ma.spawn_agent(name, persona, task)
    except Exception as e:
        return f"[error] {e}"


@register_tool("send_to_agent", "Send a message to another agent", {
    "agent_name": {"type": "string", "description": "Target agent name"},
    "message": {"type": "string", "description": "Message to send"},
})
def tool_send_to_agent(agent_name: str, message: str, **kwargs) -> str:
    from core.multi_agent import PFMultiAgent
    try:
        ma = PFMultiAgent("main", os.path.expanduser("~/.permafrost"))
        ma.send(agent_name, message)
        return f"Message sent to agent '{agent_name}'"
    except Exception as e:
        return f"[error] {e}"


@register_tool("list_agents", "List all spawned AI agents", {})
def tool_list_agents(**kwargs) -> str:
    from core.multi_agent import PFMultiAgent
    try:
        ma = PFMultiAgent("main", os.path.expanduser("~/.permafrost"))
        agents = ma.list_agents()
        if not agents:
            return "No agents spawned yet."
        lines = []
        for a in agents:
            lines.append(f"  {a.get('name','?')} — {a.get('persona','')[:60]} (created: {a.get('created_at','?')[:10]})")
        return f"Agents ({len(agents)}):\n" + "\n".join(lines)
    except Exception as e:
        return f"[error] {e}"


@register_tool("read_agent_outbox", "Read messages from an agent's outbox", {
    "agent_name": {"type": "string", "description": "Agent name to read from"},
})
def tool_read_agent_outbox(agent_name: str, **kwargs) -> str:
    from core.multi_agent import PFMultiAgent
    try:
        ma = PFMultiAgent("main", os.path.expanduser("~/.permafrost"))
        msgs = ma.check_inbox()  # Check our own inbox for messages from that agent
        relevant = [m for m in msgs if m.get("from") == agent_name]
        if not relevant:
            return f"No messages from agent '{agent_name}'."
        lines = []
        for m in relevant[-5:]:
            lines.append(f"  [{m.get('timestamp','')[:16]}] {m.get('message','')[:200]}")
        ma.mark_read()
        return f"Messages from {agent_name}:\n" + "\n".join(lines)
    except Exception as e:
        return f"[error] {e}"


# ── Self-Evolution (Rules + Tools) ────────────────────────────

@register_tool("update_rules", "Add or update a self-learned rule (persists across restarts)", {
    "rule": {"type": "string", "description": "The rule to add (e.g. 'Always confirm before deleting files')"},
    "category": {"type": "string", "description": "Category: learned, correction, preference, pitfall"},
})
def tool_update_rules(rule: str, category: str = "learned", **kwargs) -> str:
    """AI adds a rule to its own my_rules.md file. These persist and are loaded every session."""
    from pathlib import Path
    from datetime import datetime
    data_dir = Path(os.path.expanduser("~/.permafrost"))
    rules_file = data_dir / "memory" / "L1" / "my_rules.md"
    try:
        content = ""
        if rules_file.exists():
            content = rules_file.read_text(encoding="utf-8")
        # Remove placeholder if present
        content = content.replace("(none yet — use update_rules tool to add)\n", "")
        # Add the new rule
        entry = f"- [{category}] {rule} ({datetime.now().strftime('%Y-%m-%d')})\n"
        if entry.strip() in content:
            return f"Rule already exists: {rule[:60]}"
        content += entry
        rules_file.write_text(content, encoding="utf-8")
        return f"Rule added [{category}]: {rule[:80]}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("read_rules", "Read your current self-learned rules", {})
def tool_read_rules(**kwargs) -> str:
    """Read AI's own custom rules file."""
    from pathlib import Path
    rules_file = Path(os.path.expanduser("~/.permafrost")) / "memory" / "L1" / "my_rules.md"
    if not rules_file.exists():
        return "No custom rules yet."
    return rules_file.read_text(encoding="utf-8")[:4000]


# ── Self-Tool-Creation ───────────────────────────────────────

@register_tool("create_tool", "Create a new custom tool that persists across restarts", {
    "name": {"type": "string", "description": "Tool name (snake_case, e.g. 'weather_check')"},
    "description": {"type": "string", "description": "What the tool does"},
    "parameters": {"type": "string", "description": "JSON string of parameters, e.g. '{\"city\": {\"type\": \"string\", \"description\": \"City name\"}}'"},
    "code": {"type": "string", "description": "Python function body (receives kwargs matching parameters, must return a string)"},
})
def tool_create_tool(name: str, description: str, parameters: str, code: str, **kwargs) -> str:
    """Create a custom tool and register it to the auto_tools plugin directory.

    The tool is immediately available and persists across brain restarts.
    """
    from pathlib import Path

    # Validate name
    if not re.match(r'^[a-z][a-z0-9_]*$', name):
        return f"[error] Tool name must be snake_case (e.g. 'my_tool'), got '{name}'"

    if name in TOOLS:
        return f"[error] Tool '{name}' already exists. Use a different name."

    # Validate parameters JSON
    try:
        params = json.loads(parameters)
        if not isinstance(params, dict):
            return "[error] Parameters must be a JSON object"
    except json.JSONDecodeError as e:
        return f"[error] Invalid parameters JSON: {e}"

    # Build the tool file
    data_dir = Path(os.path.expanduser("~/.permafrost"))
    tools_dir = data_dir / "plugins" / "auto_tools"
    tools_dir.mkdir(parents=True, exist_ok=True)

    # Create plugin manifest if not exists
    manifest = tools_dir / "plugin.json"
    if not manifest.exists():
        manifest.write_text(json.dumps({
            "name": "auto_tools",
            "version": "1.0.0",
            "description": "Auto-generated tools created by AI",
            "author": "Permafrost AI",
        }, indent=2), encoding="utf-8")

    # Write the tool file
    tool_file = tools_dir / f"{name}.py"
    param_names = list(params.keys())
    func_params = ", ".join([f"{p}=None" for p in param_names]) + (", **kwargs" if param_names else "**kwargs")

    tool_code = f'''"""Auto-generated tool: {name}"""
import json, os, re, subprocess
from core.tools import register_tool

@register_tool("{name}", """{description}""", {json.dumps(params, ensure_ascii=False)})
def tool_{name}({func_params}):
{chr(10).join("    " + line for line in code.strip().split(chr(10)))}
'''

    tool_file.write_text(tool_code, encoding="utf-8")

    # Update __init__.py to import all tool files
    init_file = tools_dir / "__init__.py"
    imports = []
    for f in tools_dir.glob("*.py"):
        if f.name != "__init__.py":
            imports.append(f"from . import {f.stem}")
    init_file.write_text("\n".join(imports) + "\n", encoding="utf-8")

    # Immediately register the tool in current session
    try:
        exec(f"from core.tools import register_tool\n"
             f"@register_tool('{name}', '''{description}''', {json.dumps(params)})\n"
             f"def tool_{name}({func_params}):\n"
             + "\n".join(f"    {line}" for line in code.strip().split("\n")),
             {"register_tool": register_tool, "json": json, "os": os, "re": re, "subprocess": subprocess})
        return f"Tool '{name}' created and registered! ({len(params)} params). File: {tool_file}"
    except Exception as e:
        return f"Tool file created at {tool_file}, but live registration failed: {e}. It will load on next restart."


# ── Memory Tools (L1-L6 layered system) ──────────────────────

@register_tool("memory_save", "Save information to L2 verified knowledge (long-term)", {
    "name": {"type": "string", "description": "Memory name/title"},
    "content": {"type": "string", "description": "Content to remember"},
    "type": {"type": "string", "description": "Type: user/feedback/project/reference"},
})
def tool_memory_save(name, content, type="reference", **kwargs):
    from smart.memory import PFMemory
    try:
        mem = PFMemory()
        mem.save_l2(name, name, type, content)
        # Auto-index to vector store for semantic search
        try:
            from smart.vector import PFVectorSearch
            vs = PFVectorSearch(str(mem.data_dir))
            vs.index_memory(f"L2:{name}", f"{name} {content}", {"layer": "L2", "type": type})
        except Exception:
            pass  # Vector indexing is optional
        return f"[L2] Saved: {name} ({type})"
    except Exception as e:
        return f"[error] {e}"


@register_tool("memory_note", "Add a short-term dynamic note to L3 (auto-expires)", {
    "key": {"type": "string", "description": "Note key/title"},
    "value": {"type": "string", "description": "Note content"},
    "type": {"type": "string", "description": "Type: context(14d)/preference(30d)/progress(7d)/insight(21d)"},
})
def tool_memory_note(key, value, type="context", **kwargs):
    from smart.memory import PFMemory
    try:
        mem = PFMemory()
        importance = int(kwargs.get("importance", 3))
        mem.add_l3(key, value, type, importance)
        # Auto-index to vector store
        try:
            from smart.vector import PFVectorSearch
            vs = PFVectorSearch(str(mem.data_dir))
            vs.index_memory(f"L3:{key}", f"{key} {value}", {"layer": "L3", "type": type})
        except Exception:
            pass
        return f"[L3] Noted: {key} ({type})"
    except Exception as e:
        return f"[error] {e}"


@register_tool("memory_search", "Search across all memory layers (L2 + L3) with semantic vector search", {
    "query": {"type": "string", "description": "Search query (supports natural language semantic search)"},
})
def tool_memory_search(query, **kwargs):
    from smart.memory import PFMemory
    mem = PFMemory()

    # Try semantic search first, fallback to keyword
    results = mem.search_semantic(query, top_k=10)
    if not results:
        results = mem.search_all(query)

    if not results:
        return "No memories found."

    lines = []
    for r in results[:10]:
        # Handle both semantic and keyword result formats
        layer = r.get("layer", r.get("metadata", {}).get("layer", "?"))
        score = r.get("score", "")
        score_str = f" (score:{score:.2f})" if isinstance(score, float) else ""

        if layer == "L2":
            name = r.get("name", r.get("text", "")[:50])
            body = r.get("body", r.get("text", ""))[:200]
            lines.append(f"[L2:{r.get('type', r.get('metadata', {}).get('type', ''))}]{score_str} {name}: {body}")
        elif layer == "L3":
            key = r.get("key", "")
            value = r.get("value", r.get("text", ""))[:200]
            lines.append(f"[L3:{r.get('type', r.get('metadata', {}).get('type', ''))}]{score_str} {key}: {value}")
        else:
            lines.append(f"[{layer}]{score_str} {r.get('text', '')[:200]}")
    return "\n".join(lines)


@register_tool("memory_list", "List all saved memories across layers", {})
def tool_memory_list(**kwargs):
    from smart.memory import PFMemory, L2_TYPES
    mem = PFMemory()
    lines = []

    # L2: Verified Knowledge
    for mtype in L2_TYPES:
        items = mem.list_l2(mtype)
        if items:
            lines.append(f"\n[L2:{mtype}]")
            for i in items:
                lines.append(f"  - {i.get('name','')}: {i.get('description','')}")

    # L3: Dynamic
    l3 = mem.list_l3()
    if l3:
        lines.append(f"\n[L3] ({len(l3)} entries)")
        for e in l3[:15]:
            lines.append(f"  - [{e.get('type','')}] {e.get('key','')}: {e.get('value','')[:80]}")

    return "\n".join(lines) if lines else "No memories saved yet."


@register_tool("memory_gc", "Run garbage collection on L3 dynamic memories (expire/promote/archive)", {})
def tool_memory_gc(**kwargs):
    from smart.memory import PFMemory
    try:
        mem = PFMemory()
        result = mem.gc()
        return f"GC complete: kept={result['kept']}, promoted={result['promoted']}, archived={result['archived']}"
    except Exception as e:
        return f"[error] {e}"


@register_tool("memory_reindex", "Rebuild vector search index from all L2+L3 memories", {})
def tool_memory_reindex(**kwargs):
    from smart.memory import PFMemory
    try:
        mem = PFMemory()
        mem.index_all_memories()
        stats = mem.get_stats()
        vectors = stats.get("vectors", 0)
        return f"Vector index rebuilt: {vectors} entries indexed (L2+L3)"
    except Exception as e:
        return f"[error] {e}"


@register_tool("memory_stats", "Show memory layer statistics (L1-L6 + vector index)", {})
def tool_memory_stats(**kwargs):
    from smart.memory import PFMemory
    mem = PFMemory()
    stats = mem.get_stats()
    lines = [f"  {layer}: {count}" for layer, count in stats.items()]
    return "Memory Stats:\n" + "\n".join(lines)


@register_tool("set_reminder", "Set a timed reminder that will notify the user at the specified time", {
    "message": {"type": "string", "description": "Reminder message to send"},
    "time": {"type": "string", "description": "Time in HH:MM format (24h)"},
    "repeat": {"type": "string", "description": "once (default), daily, or weekly"},
})
def tool_set_reminder(message, time, repeat="once", **kwargs):
    """Create a scheduled reminder. Scheduler will fire it at the specified time."""
    import uuid
    from datetime import datetime
    from pathlib import Path

    data_dir = Path(os.path.expanduser("~/.permafrost"))
    reminder_file = data_dir / "reminders.json"

    # Validate time format
    try:
        datetime.strptime(time, "%H:%M")
    except ValueError:
        return f"[error] Invalid time format '{time}'. Use HH:MM (e.g. 22:00)"

    if repeat not in ("once", "daily", "weekly"):
        repeat = "once"

    # Load existing reminders
    reminders = []
    if reminder_file.exists():
        try:
            reminders = json.loads(reminder_file.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            reminders = []

    reminder_id = f"rem-{uuid.uuid4().hex[:8]}"
    reminders.append({
        "id": reminder_id,
        "message": message,
        "time": time,
        "repeat": repeat,
        "enabled": True,
        "created": datetime.now().isoformat(),
    })

    reminder_file.write_text(
        json.dumps(reminders, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    return f"Reminder set: '{message}' at {time} ({repeat}). ID: {reminder_id}"


@register_tool("list_reminders", "List all active reminders", {})
def tool_list_reminders(**kwargs):
    """List all scheduled reminders."""
    from pathlib import Path
    data_dir = Path(os.path.expanduser("~/.permafrost"))
    reminder_file = data_dir / "reminders.json"
    if not reminder_file.exists():
        return "No reminders set."
    try:
        reminders = json.loads(reminder_file.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return "No reminders set."
    if not reminders:
        return "No reminders set."
    lines = []
    for r in reminders:
        status = "ON" if r.get("enabled", True) else "OFF"
        lines.append(f"  [{status}] {r.get('time','')} ({r.get('repeat','once')}) — {r.get('message','')}")
    return f"Reminders ({len(reminders)}):\n" + "\n".join(lines)


@register_tool("delete_reminder", "Delete a reminder by ID", {
    "reminder_id": {"type": "string", "description": "Reminder ID to delete"},
})
def tool_delete_reminder(reminder_id, **kwargs):
    """Delete a scheduled reminder."""
    from pathlib import Path
    data_dir = Path(os.path.expanduser("~/.permafrost"))
    reminder_file = data_dir / "reminders.json"
    if not reminder_file.exists():
        return "No reminders found."
    try:
        reminders = json.loads(reminder_file.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return "No reminders found."
    before = len(reminders)
    reminders = [r for r in reminders if r.get("id") != reminder_id]
    if len(reminders) == before:
        return f"Reminder '{reminder_id}' not found."
    reminder_file.write_text(
        json.dumps(reminders, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    return f"Reminder '{reminder_id}' deleted."


# ── Tool Executor ─────────────────────────────────────────────

def execute_tool(name: str, args: dict, security=None) -> str:
    """Execute a registered tool by name.

    Args:
        name: Tool name (must be in TOOLS registry).
        args: Keyword arguments to pass to the tool function.
        security: Optional PFSecurity instance for authorization.

    Returns:
        Tool output string (always returns a string, never raises).
    """
    if name not in TOOLS:
        return f"[error] Unknown tool: {name}"

    # Security check via PFSecurity.authorize_tool()
    if security:
        allowed, reason = security.authorize_tool(name, args)
        if not allowed:
            log.warning(f"Tool '{name}' blocked: {reason}")
            return f"[blocked] Tool '{name}' not allowed: {reason}"

    try:
        return TOOLS[name]["function"](**args)
    except Exception as e:
        log.error(f"Tool '{name}' execution failed: {e}")
        return f"[error] Tool execution failed: {e}"


# ── Schema Export ─────────────────────────────────────────────

def get_tool_schemas() -> list[dict]:
    """Return tool schemas in OpenAI function-calling format.

    Useful for providers that support native function calling.
    """
    return [
        {
            "type": "function",
            "function": {
                "name": name,
                "description": info["description"],
                "parameters": {
                    "type": "object",
                    "properties": info["parameters"],
                },
            },
        }
        for name, info in TOOLS.items()
    ]


def get_tool_prompt() -> str:
    """Generate a text prompt describing available tools.

    Used for prompt-injection-based tool use (works with any LLM).
    """
    lines = [
        "You have access to these tools to interact with the computer:",
        "",
    ]
    for idx, (name, info) in enumerate(TOOLS.items(), 1):
        param_parts = []
        for pname in info["parameters"]:
            param_parts.append(pname)
        params_str = ", ".join(param_parts)
        lines.append(f"{idx}. {name}({params_str}) — {info['description']}")

    lines.extend([
        "",
        "IMPORTANT: To use a tool, you MUST use EXACTLY this format:",
        '[TOOL_CALL]{"name": "tool_name", "args": {"key": "value"}}[/TOOL_CALL]',
        "",
        "Example — user asks to remember something:",
        'Sure, I\'ll save that.',
        '[TOOL_CALL]{"name": "memory_note", "args": {"key": "preference", "value": "likes coffee", "type": "preference"}}[/TOOL_CALL]',
        "",
        "Example — user asks to search memory:",
        '[TOOL_CALL]{"name": "memory_search", "args": {"query": "coffee"}}[/TOOL_CALL]',
        "",
        "Rules:",
        "- Tag MUST be [TOOL_CALL] and [/TOOL_CALL]. No other tag names.",
        "- JSON must use double quotes.",
        "- After each tool call, you receive the result and can respond naturally.",
        "- Use tools for files, commands, memory, and information retrieval.",
    ])
    return "\n".join(lines)


# ── Tool Call Normalizer + Parser ─────────────────────────────

_TOOL_CALL_PATTERN = re.compile(
    r"\[TOOL_CALL\]\s*(\{.*?\})\s*\[/TOOL_CALL\]",
    re.DOTALL,
)

# Catch-all: any variant AI models might invent
_TOOL_CALL_ANY = re.compile(
    r"\[(?:TOOL_CALL|TOOL_CODE|tool_call|tool_code|Tool_Call|ToolCall|toolcall)\]"
    r"\s*(\{.*?\})\s*"
    r"\[/(?:TOOL_CALL|TOOL_CODE|tool_call|tool_code|Tool_Call|ToolCall|toolcall)\]",
    re.DOTALL,
)

# Even more aggressive: ```tool_call blocks (GPT-style)
_TOOL_CALL_BACKTICK = re.compile(
    r"```(?:tool_call|json)?\s*(\{[^`]*?\"name\"\s*:[^`]*?\})\s*```",
    re.DOTALL,
)


def normalize_tool_calls(text: str) -> str:
    """Normalize ANY tool call format into standard [TOOL_CALL]...[/TOOL_CALL].

    Handles:
      - [TOOL_CODE]...[/TOOL_CODE] (Gemini)
      - [tool_call]...[/tool_call] (lowercase variants)
      - ```json {...} ``` (GPT-style code blocks with tool JSON)
      - Any other bracket variant

    This is the key to model-agnostic tool use: don't trust the model
    to use the right format, just fix whatever it outputs.
    """
    # Already standard? Return as-is
    if _TOOL_CALL_PATTERN.search(text):
        return text

    # Try bracket variants
    normalized = _TOOL_CALL_ANY.sub(
        lambda m: f"[TOOL_CALL]{m.group(1)}[/TOOL_CALL]", text
    )
    if _TOOL_CALL_PATTERN.search(normalized):
        return normalized

    # Try backtick code blocks containing tool JSON
    for match in _TOOL_CALL_BACKTICK.finditer(text):
        raw = match.group(1).strip()
        try:
            data = json.loads(raw)
            if "name" in data:
                replacement = f'[TOOL_CALL]{raw}[/TOOL_CALL]'
                normalized = text.replace(match.group(0), replacement)
                return normalized
        except (json.JSONDecodeError, ValueError):
            pass

    return text


def parse_tool_calls(text: str) -> list[dict]:
    """Parse [TOOL_CALL]...[/TOOL_CALL] blocks from AI response.

    Returns list of {"name": str, "args": dict} dicts.
    Invalid JSON blocks are skipped with a warning.
    """
    calls = []
    # Try standard format first, then variants (TOOL_CODE, etc.)
    matches = list(_TOOL_CALL_PATTERN.finditer(text))
    if not matches:
        matches = list(_TOOL_CALL_VARIANTS.finditer(text))
    for match in matches:
        raw = match.group(1)
        try:
            data = json.loads(raw)
            name = data.get("name", "")
            args = data.get("args", {})
            if name:
                calls.append({"name": name, "args": args})
            else:
                log.warning(f"Tool call missing 'name': {raw[:100]}")
        except json.JSONDecodeError as e:
            log.warning(f"Invalid tool call JSON: {e} — {raw[:100]}")
    return calls


def has_tool_calls(text: str) -> bool:
    """Check if text contains any tool call blocks (standard or variant)."""
    return bool(_TOOL_CALL_PATTERN.search(text) or _TOOL_CALL_VARIANTS.search(text))


def strip_tool_calls(text: str) -> str:
    """Remove ALL tool call blocks from text — aggressive cleanup.

    Catches [TOOL_CALL], [TOOL_CODE], [tool_call], backtick blocks,
    and any remaining JSON-looking tool invocations.
    """
    result = _TOOL_CALL_PATTERN.sub("", text)
    result = _TOOL_CALL_VARIANTS.sub("", result)
    result = _TOOL_CALL_BACKTICK.sub("", result)
    # Last resort: catch any remaining [ANYTHING]{json}[/ANYTHING] patterns
    result = re.sub(
        r'\[/?(?:TOOL_CALL|TOOL_CODE|tool_call|tool_code)\]',
        '', result,
    )
    # Clean up any orphaned JSON tool objects left behind
    result = re.sub(
        r'\{"name"\s*:\s*"[a-z_]+"[^}]*"args"\s*:\s*\{[^}]*\}\s*\}',
        '', result,
    )
    # Remove excessive whitespace from cleanup
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip()
